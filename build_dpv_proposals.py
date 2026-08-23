from __future__ import annotations

import math
import os
import textwrap
import urllib.request
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "deliverables"
ASSETS = ROOT / "doc_assets"
OUT.mkdir(exist_ok=True)
ASSETS.mkdir(exist_ok=True)

RED = "C62828"
DARK_RED = "8E1B1B"
NAVY = "18324A"
BLUE = "2E6F9E"
LIGHT_BLUE = "EAF2F8"
LIGHT_RED = "FBEDED"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "D7DDE3"
GRAY = "5D6873"
INK = "20262D"
GREEN = "2E7D5B"
LIGHT_GREEN = "E9F5EF"
AMBER = "C88719"
LIGHT_AMBER = "FFF4DA"
WHITE = "FFFFFF"

FONT_CJK = Path(r"C:\Windows\Fonts\msyh.ttc")
FONT_CJK_BOLD = Path(r"C:\Windows\Fonts\msyhbd.ttc")
FONT_EN = Path(r"C:\Windows\Fonts\arial.ttf")
FONT_EN_BOLD = Path(r"C:\Windows\Fonts\arialbd.ttf")


def pil_font(lang: str, size: int, bold: bool = False):
    path = FONT_CJK_BOLD if lang == "zh" and bold else FONT_CJK if lang == "zh" else FONT_EN_BOLD if bold else FONT_EN
    return ImageFont.truetype(str(path), size=size)


def wrap_label(draw: ImageDraw.ImageDraw, text: str, font, max_width: int):
    words = list(text) if any("\u4e00" <= c <= "\u9fff" for c in text) else text.split()
    lines, current = [], ""
    sep = "" if words and len(words[0]) == 1 and any("\u4e00" <= c <= "\u9fff" for c in text) else " "
    for word in words:
        trial = word if not current else current + sep + word
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def rounded_box(draw, box, fill, outline, text, lang, font_size=38, radius=22, width=4, bold=False):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)
    font = pil_font(lang, font_size, bold)
    lines = wrap_label(draw, text, font, int((x2 - x1) * 0.86))
    line_h = font_size + 10
    y = (y1 + y2 - line_h * len(lines)) / 2
    for line in lines:
        b = draw.textbbox((0, 0), line, font=font)
        x = (x1 + x2 - (b[2] - b[0])) / 2
        draw.text((x, y), line, font=font, fill="#20262D")
        y += line_h


def arrow(draw, start, end, color="#5D6873", width=7, dashed=False):
    x1, y1 = start
    x2, y2 = end
    if dashed:
        segs = 14
        for i in range(segs):
            if i % 2 == 0:
                a = i / segs
                b = min((i + 1) / segs, 1)
                draw.line((x1 + (x2-x1)*a, y1 + (y2-y1)*a, x1 + (x2-x1)*b, y1 + (y2-y1)*b), fill=color, width=width)
    else:
        draw.line((x1, y1, x2, y2), fill=color, width=width)
    angle = math.atan2(y2-y1, x2-x1)
    length = 24
    for off in (2.55, -2.55):
        draw.line((x2, y2, x2 + length*math.cos(angle+off), y2 + length*math.sin(angle+off)), fill=color, width=width)


def title_on_canvas(draw, text, lang, width, subtitle=None):
    f = pil_font(lang, 58, True)
    b = draw.textbbox((0, 0), text, font=f)
    draw.text(((width-(b[2]-b[0]))/2, 38), text, font=f, fill="#18324A")
    draw.rectangle((width*0.38, 115, width*0.62, 124), fill="#C62828")
    if subtitle:
        sf = pil_font(lang, 27)
        sb = draw.textbbox((0, 0), subtitle, font=sf)
        draw.text(((width-(sb[2]-sb[0]))/2, 142), subtitle, font=sf, fill="#5D6873")


def save_image(img: Image.Image, name: str):
    path = ASSETS / name
    img.save(path, dpi=(300, 300), optimize=True)
    return path


def source_map(lang: str):
    zh = lang == "zh"
    W, H = 2400, 1450
    img = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(img)
    title_on_canvas(d, "海外多平台客户获取来源图" if zh else "Multi-Platform Overseas Lead Sources", lang, W,
                    "公开企业信号统一进入潜客识别与验证流程" if zh else "Public business signals consolidated into one qualification pipeline")
    left = [
        "Google搜索与地图" if zh else "Google Search & Maps",
        "LinkedIn进口商与批发商" if zh else "LinkedIn Importers & Wholesalers",
        "Facebook/Instagram企业页" if zh else "Facebook / Instagram Business Pages",
        "TikTok/YouTube行业信号" if zh else "TikTok / YouTube Industry Signals",
        "B2B采购平台" if zh else "B2B Marketplaces",
    ]
    right = [
        "进出口贸易数据库" if zh else "Import / Export Data",
        "批发商与分销商目录" if zh else "Wholesaler & Distributor Directories",
        "海外展会参展商名单" if zh else "Trade Show Exhibitor Lists",
        "连锁渠道供应商线索" if zh else "Chain-Store Supplier Leads",
        "自有B2B询盘" if zh else "Owned B2B Inquiries",
    ]
    ys = [270, 470, 670, 870, 1070]
    for x, items, fill in [(80, left, "#EAF2F8"), (1780, right, "#FBEDED")]:
        for y, label in zip(ys, items):
            rounded_box(d, (x, y, x+540, y+135), fill, "#2E6F9E" if x < 1000 else "#C62828", label, lang, 34, bold=True)
            arrow(d, (x+540 if x < 1000 else x, y+67), (930 if x < 1000 else 1470, 665), color="#7A8793", width=5)
    rounded_box(d, (930, 515, 1470, 815), "#FFF4DA", "#C88719", "统一潜客数据库" if zh else "Unified Lead Database", lang, 44, bold=True)
    rounded_box(d, (905, 930, 1495, 1075), "#F3F5F7", "#5D6873", "清洗 · 去重 · 验证" if zh else "Clean · Deduplicate · Verify", lang, 38, bold=True)
    rounded_box(d, (905, 1165, 1495, 1310), "#E9F5EF", "#2E7D5B", "分类 · 评分 · 分配" if zh else "Classify · Score · Route", lang, 38, bold=True)
    arrow(d, (1200, 815), (1200, 930), color="#C88719")
    arrow(d, (1200, 1075), (1200, 1165), color="#2E7D5B")
    return save_image(img, f"source_map_{lang}.png")


def acquisition_flow(lang: str):
    zh = lang == "zh"
    W, H = 2800, 1700
    img = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(img)
    title_on_canvas(d, "AI海外自动获客端到端流程" if zh else "End-to-End AI-Powered Lead Generation Workflow", lang, W,
                    "从目标定义到成交反馈的可追溯闭环" if zh else "A traceable closed loop from market definition to conversion feedback")
    rows = [
        (["确定市场与批发商画像", "生成多语言关键词", "多平台搜索", "提取企业资料", "清洗去重"] if zh else
         ["Define Market & Wholesaler ICP", "Generate Search Terms", "Search Platforms", "Extract Company Data", "Clean & Deduplicate"]),
        (["补全采购联系人", "判断进口批发资质", "商品匹配", "客户评分", "生成开发话术"] if zh else
         ["Enrich Buyer Contacts", "Qualify Import / Wholesale Fit", "Match Products", "Score Leads", "Draft Outreach"]),
        (["人工审批", "多渠道触达", "回复分类", "销售跟进", "成交反馈"] if zh else
         ["Human Approval", "Multichannel Outreach", "Classify Replies", "Sales Follow-up", "Conversion Feedback"]),
    ]
    xs = [100, 635, 1170, 1705, 2240]
    ys = [270, 690, 1110]
    fills = ["#EAF2F8", "#F3F5F7", "#FBEDED"]
    outlines = ["#2E6F9E", "#5D6873", "#C62828"]
    for ri, labels in enumerate(rows):
        for ci, label in enumerate(labels):
            rounded_box(d, (xs[ci], ys[ri], xs[ci]+450, ys[ri]+150), fills[ri], outlines[ri], label, lang, 31, bold=True)
            if ci < 4:
                arrow(d, (xs[ci]+450, ys[ri]+75), (xs[ci+1]-18, ys[ri]+75), color=outlines[ri], width=6)
        if ri < 2:
            arrow(d, (2465, ys[ri]+150), (2465, ys[ri+1]-20), color="#7A8793", width=6)
            arrow(d, (2465, ys[ri+1]-20), (100, ys[ri+1]-20), color="#7A8793", width=6)
            arrow(d, (100, ys[ri+1]-20), (100, ys[ri+1]), color="#7A8793", width=6)
    # Branches and feedback
    rounded_box(d, (720, 1480, 1220, 1605), "#FBEDED", "#C62828", "无效/退订：停止" if zh else "Invalid / Opt-out: Stop", lang, 29, bold=True)
    rounded_box(d, (1580, 1480, 2130, 1605), "#E9F5EF", "#2E7D5B", "高意向：转销售" if zh else "High Intent: Sales Handoff", lang, 29, bold=True)
    arrow(d, (860, 1260), (970, 1480), color="#C62828", width=5, dashed=True)
    arrow(d, (1920, 1260), (1850, 1480), color="#2E7D5B", width=5, dashed=True)
    # feedback loop from conversion back to scoring
    d.line((2690, 1185, 2730, 1185, 2730, 830, 1930, 830), fill="#C88719", width=6)
    arrow(d, (2730, 830), (1930, 830), color="#C88719", width=6)
    f = pil_font(lang, 26, True)
    feedback = "成交数据优化评分" if zh else "Conversion data improves scoring"
    d.text((2100, 785), feedback, font=f, fill="#8A6112")
    return save_image(img, f"acquisition_flow_{lang}.png")


def architecture(lang: str):
    zh = lang == "zh"
    W, H = 3000, 1650
    img = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(img)
    title_on_canvas(d, "AI自动获客系统技术架构" if zh else "AI-Powered Lead Generation System Architecture", lang, W,
                    "数据来源、工作流、智能能力、业务系统与触达渠道分层解耦" if zh else "Decoupled layers for sources, orchestration, intelligence, operations and outreach")
    cols = [
        (80, 500, "数据来源层" if zh else "Data Sources", "#EAF2F8", "#2E6F9E",
         ["搜索/地图", "B2B社媒信号", "B2B/贸易数据", "展会/行业目录"] if zh else ["Search / Maps", "B2B Social Signals", "B2B / Trade Data", "Events / Directories"]),
        (600, 980, "工作流层" if zh else "Orchestration", "#FFF4DA", "#C88719",
         ["n8n调度", "API连接", "审批与告警"] if zh else ["n8n Scheduling", "API Connectors", "Approval & Alerts"]),
        (1080, 1510, "AI能力层" if zh else "AI Services", "#F3F5F7", "#5D6873",
         ["企业分类", "商品匹配", "多语言文案", "回复识别"] if zh else ["Company Classification", "Product Matching", "Multilingual Copy", "Reply Intent"]),
        (1610, 2040, "数据与业务层" if zh else "Data & Operations", "#E9F5EF", "#2E7D5B",
         ["PostgreSQL", "CRM/NocoDB", "商品知识库", "审计日志"] if zh else ["PostgreSQL", "CRM / NocoDB", "Product Knowledge", "Audit Log"]),
        (2140, 2550, "触达层" if zh else "Outreach", "#FBEDED", "#C62828",
         ["企业邮箱", "LinkedIn任务", "WhatsApp", "销售分配"] if zh else ["Business Email", "LinkedIn Tasks", "WhatsApp", "Sales Routing"]),
        (2650, 2940, "管理层" if zh else "Management", "#EEF0F3", "#5D6873",
         ["渠道看板", "转化报表", "异常监控"] if zh else ["Channel Dashboard", "Conversion Reports", "Exception Monitor"]),
    ]
    for x1, x2, heading, fill, stroke, items in cols:
        d.rounded_rectangle((x1, 235, x2, 1470), radius=28, fill=fill, outline=stroke, width=5)
        hf = pil_font(lang, 34, True)
        hb = d.textbbox((0, 0), heading, font=hf)
        d.text(((x1+x2-(hb[2]-hb[0]))/2, 275), heading, font=hf, fill="#18324A")
        y = 405
        for item in items:
            rounded_box(d, (x1+38, y, x2-38, y+155), "#FFFFFF", stroke, item, lang, 28, radius=18, width=3, bold=True)
            y += 225
    for i in range(len(cols)-1):
        arrow(d, (cols[i][1], 855), (cols[i+1][0]-15, 855), color="#7A8793", width=7)
    # governance line
    d.line((760, 1530, 2720, 1530), fill="#C62828", width=5)
    gf = pil_font(lang, 26, True)
    gov = "统一权限、退订、审计与数据留存控制" if zh else "Unified access, opt-out, audit and retention controls"
    gb = d.textbbox((0, 0), gov, font=gf)
    d.text(((W-(gb[2]-gb[0]))/2, 1550), gov, font=gf, fill="#8E1B1B")
    return save_image(img, f"architecture_{lang}.png")


def scoring(lang: str):
    zh = lang == "zh"
    W, H = 2400, 1600
    img = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(img)
    title_on_canvas(d, "潜客评分与客户分层" if zh else "Lead Scoring and Qualification Tiers", lang, W,
                    "100分规则模型支持可解释的销售优先级" if zh else "A transparent 100-point model for sales prioritization")
    factors = [
        ("商品匹配" if zh else "Product Fit", 20),
        ("市场匹配" if zh else "Market Fit", 15),
        ("进口批发商匹配" if zh else "Importer / Wholesaler Fit", 15),
        ("进口/批发证据" if zh else "Import / Wholesale Evidence", 15),
        ("公司规模" if zh else "Company Scale", 10),
        ("近期采购信号" if zh else "Recent Buying Signal", 10),
        ("采购负责人" if zh else "Decision Maker", 10),
        ("联系方式有效" if zh else "Contact Validity", 5),
    ]
    lf = pil_font(lang, 29, True)
    vf = pil_font(lang, 28, True)
    y = 270
    for label, score in factors:
        d.text((90, y+7), label, font=lf, fill="#20262D")
        d.rounded_rectangle((510, y, 510+score*25, y+55), radius=12, fill="#2E6F9E", outline="#18324A", width=2)
        d.text((1040, y+7), str(score), font=vf, fill="#18324A")
        y += 125
    # funnel on right
    stages = [
        ("候选进口商/批发商" if zh else "Candidate Importers / Wholesalers", 1200, "#EAF2F8"),
        ("验证与匹配" if zh else "Verified and Matched", 980, "#DDEBF5"),
        ("A级 75–100" if zh else "Tier A 75–100", 760, "#FFF4DA"),
        ("85+ 自动化候选" if zh else "85+ Automation Eligible", 540, "#E9F5EF"),
    ]
    center = 1790
    y = 300
    for label, width, fill in stages:
        h = 230
        pts = [(center-width/2, y), (center+width/2, y), (center+width*0.42, y+h), (center-width*0.42, y+h)]
        d.polygon(pts, fill=fill, outline="#5D6873")
        font = pil_font(lang, 34, True)
        b = d.textbbox((0, 0), label, font=font)
        d.text((center-(b[2]-b[0])/2, y+85), label, font=font, fill="#20262D")
        y += h+25
    rounded_box(d, (1320, 1390, 2260, 1515), "#F3F5F7", "#5D6873", "B级培育 · C级观察/淘汰" if zh else "Tier B Nurture · Tier C Observe / Exclude", lang, 28, bold=True)
    return save_image(img, f"scoring_{lang}.png")


def timeline(lang: str):
    zh = lang == "zh"
    W, H = 3000, 1600
    img = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(img)
    title_on_canvas(d, "八周项目实施进度" if zh else "Eight-Week Implementation Roadmap", lang, W,
                    "第2周演示、第4周MVP、第5周试运行、第8周验收" if zh else "Demo in Week 2, MVP in Week 4, pilot in Week 5 and acceptance in Week 8")
    left = 690
    top = 300
    col_w = 260
    row_h = 120
    # header
    hf = pil_font(lang, 29, True)
    for w in range(8):
        x = left + w*col_w
        d.rectangle((x, top, x+col_w, top+70), fill="#18324A", outline="white", width=2)
        label = f"第{w+1}周" if zh else f"Week {w+1}"
        b = d.textbbox((0,0), label, font=hf)
        d.text((x+(col_w-(b[2]-b[0]))/2, top+16), label, font=hf, fill="white")
    tasks = [
        ("需求与数据盘点" if zh else "Requirements & Data", 1, 1, "#5D6873"),
        ("快速原型" if zh else "Rapid Prototype", 2, 2, "#2E6F9E"),
        ("多平台采集MVP" if zh else "Multi-Source MVP", 3, 4, "#2E6F9E"),
        ("分类/匹配/话术" if zh else "Classify / Match / Copy", 4, 5, "#C88719"),
        ("审批与真实触达" if zh else "Approval & Live Outreach", 5, 5, "#C62828"),
        ("B2B信号与看板" if zh else "B2B Signals & Dashboard", 6, 6, "#7B5EA7"),
        ("优化与市场复制" if zh else "Optimize & Replicate", 7, 7, "#2E7D5B"),
        ("测试/培训/验收" if zh else "Test / Train / Accept", 8, 8, "#18324A"),
    ]
    tf = pil_font(lang, 30, True)
    for i, (label, start, endw, color) in enumerate(tasks):
        y = top+95+i*row_h
        d.text((70, y+24), label, font=tf, fill="#20262D")
        for w in range(8):
            x = left+w*col_w
            d.rectangle((x, y, x+col_w, y+85), fill="#FAFBFC", outline="#D7DDE3", width=2)
        x1 = left+(start-1)*col_w+10
        x2 = left+endw*col_w-10
        d.rounded_rectangle((x1, y+13, x2, y+72), radius=18, fill=color)
    # milestones
    milestones = [(2, "原型演示" if zh else "Demo"), (4, "MVP完成" if zh else "MVP"), (5, "试运行" if zh else "Pilot"), (8, "首期验收" if zh else "Acceptance")]
    mf = pil_font(lang, 25, True)
    for week, label in milestones:
        x = left+(week-0.5)*col_w
        y = 1375
        d.polygon([(x,y-24),(x+24,y),(x,y+24),(x-24,y)], fill="#C62828")
        b = d.textbbox((0,0), label, font=mf)
        d.text((x-(b[2]-b[0])/2, y+35), label, font=mf, fill="#8E1B1B")
    return save_image(img, f"timeline_{lang}.png")


def download_logo():
    path = ASSETS / "dpv-logo.jpeg"
    if not path.exists():
        try:
            urllib.request.urlretrieve("https://dpvinternational.com/images/dpv-logo.jpeg", path)
        except Exception:
            return None
    return path if path.exists() else None


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths_dxa)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.first_child_found_in("w:tblInd")
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(indent))
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.first_child_found_in("w:tcW")
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths_dxa[min(i, len(widths_dxa)-1)]))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, lang, size=None, bold=None, color=None, italic=None):
    name = "Microsoft YaHei" if lang == "zh" else "Arial"
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.extend([fldChar1, instrText, fldChar2])


def configure_section(section, landscape=False):
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.8)
    section.header_distance = Inches(0.36)
    section.footer_distance = Inches(0.36)


def setup_styles(doc, lang):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei" if lang == "zh" else "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei" if lang == "zh" else "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for sty_name, size, color, before, after in [
        ("Title", 28, NAVY, 0, 8),
        ("Subtitle", 13, GRAY, 0, 10),
        ("Heading 1", 16, RED, 16, 8),
        ("Heading 2", 13, NAVY, 12, 6),
        ("Heading 3", 11.5, BLUE, 8, 4),
    ]:
        st = styles[sty_name]
        st.font.name = "Microsoft YaHei" if lang == "zh" else "Arial"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei" if lang == "zh" else "Arial")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    for list_name in ("List Bullet", "List Number"):
        st = styles[list_name]
        st.font.name = "Microsoft YaHei" if lang == "zh" else "Arial"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei" if lang == "zh" else "Arial")
        st.font.size = Pt(10.5)
        st.paragraph_format.left_indent = Inches(0.375)
        st.paragraph_format.first_line_indent = Inches(-0.194)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.208
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap = styles["Caption"]
    cap.font.name = "Microsoft YaHei" if lang == "zh" else "Arial"
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei" if lang == "zh" else "Arial")
    cap.font.size = Pt(9)
    cap.font.bold = True
    cap.font.color.rgb = RGBColor.from_string(NAVY)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(4)
    cap.paragraph_format.keep_with_next = True


def set_headers_footers(doc, lang):
    left = "DPV International | AI海外自动获客系统建设方案" if lang == "zh" else "DPV International | AI-Powered Overseas Lead Generation System"
    right = "内部方案" if lang == "zh" else "Internal Proposal"
    for idx, section in enumerate(doc.sections):
        configure_section(section, section.orientation == WD_ORIENT.LANDSCAPE)
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        header = section.header
        p = header.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(left)
        set_run_font(r, lang, 8.5, True, GRAY)
        r2 = p.add_run("    |    " + right)
        set_run_font(r2, lang, 8.5, False, RED)
        p.paragraph_format.space_after = Pt(3)
        # bottom rule
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "3")
        bottom.set(qn("w:color"), MID_GRAY)
        pBdr.append(bottom)
        pPr.append(pBdr)
        fp = section.footer.paragraphs[0]
        fp.clear()
        add_page_field(fp)
        for run in fp.runs:
            set_run_font(run, lang, 8.5, False, GRAY)
    # cover page has no header/footer
    first = doc.sections[0]
    first.different_first_page_header_footer = True
    first.first_page_header.paragraphs[0].clear()
    first.first_page_footer.paragraphs[0].clear()


def add_para(doc, text, lang, bold=False, color=None, size=None, align=None, after=7, italic=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_run_font(r, lang, size, bold, color, italic)
    return p


def add_bullets(doc, items, lang, size=10.5, after=4):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, lang, size, False, INK)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(after)


def add_numbered(doc, items, lang):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r, lang, 10.5, False, INK)


def add_heading(doc, text, level, lang):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_run_font(r, lang, None, True, RED if level == 1 else NAVY if level == 2 else BLUE)
    return p


def add_table(doc, headers, rows, widths, lang, font_size=8.6):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        set_cell_shading(hdr.cells[i], NAVY)
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, lang, font_size, True, WHITE)
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if ri % 2 == 1:
                set_cell_shading(cells[i], "FAFBFC")
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 or len(str(val)) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run_font(r, lang, font_size, False, INK)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, text, lang, fill=LIGHT_RED, accent=RED):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, lang, 10.5, True, accent)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(doc, image_path, caption, description, lang, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(image_path), width=Inches(width))
    try:
        shape._inline.docPr.set("descr", caption)
    except Exception:
        pass
    cp = doc.add_paragraph(style="Caption")
    cp.add_run(caption)
    dp = add_para(doc, description, lang, size=9.2, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    dp.paragraph_format.keep_with_next = False


def add_landscape_figure(doc, title, image_path, caption, description, lang):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(sec, True)
    add_heading(doc, title, 1, lang)
    add_figure(doc, image_path, caption, description, lang, width=9.35)
    sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(sec2, False)


def cover(doc, lang, logo):
    zh = lang == "zh"
    add_para(doc, "DPV INTERNATIONAL", lang, bold=True, color=RED, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    if logo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shape = p.add_run().add_picture(str(logo), width=Inches(1.05))
        shape._inline.docPr.set("descr", "DPV International corporate logo")
        p.paragraph_format.space_after = Pt(28)
    else:
        add_para(doc, "DPV", lang, bold=True, color=RED, size=22, align=WD_ALIGN_PARAGRAPH.CENTER, after=28)
    title = "AI海外自动获客系统建设方案" if zh else "AI-Powered Overseas Lead Generation System"
    subtitle = "面向海外进口型批发商的B2B潜客发现、筛选与销售转化闭环" if zh else "B2B Lead Discovery, Qualification and Conversion for Overseas Importer-Wholesalers"
    add_para(doc, title, lang, bold=True, color=NAVY, size=27 if zh else 25, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_para(doc, subtitle, lang, color=GRAY, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, after=44)
    # metric strip
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [2340]*4, indent=0)
    metrics = [("8周" if zh else "8 Weeks", "首期建设" if zh else "Initial Delivery"),
               ("第2周" if zh else "Week 2", "原型演示" if zh else "Prototype Demo"),
               ("第5周" if zh else "Week 5", "真实试运行" if zh else "Live Pilot"),
               ("第8周" if zh else "Week 8", "项目验收" if zh else "Acceptance")]
    for cell, (big, small) in zip(table.rows[0].cells, metrics):
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(big + "\n")
        set_run_font(r, lang, 14, True, RED)
        r2 = p.add_run(small)
        set_run_font(r2, lang, 8.5, False, GRAY)
    add_para(doc, "", lang, after=35)
    meta = [
        ("编制对象" if zh else "Prepared for", "DPV International管理层及业务部门" if zh else "DPV International Management and Business Teams"),
        ("方案版本" if zh else "Version", "V1.0"),
        ("方案日期" if zh else "Date", "2026年8月" if zh else "August 2026"),
        ("文件属性" if zh else "Classification", "内部方案" if zh else "Internal Proposal"),
    ]
    add_table(doc, ["项目" if zh else "Item", "内容" if zh else "Details"], meta, [2100, 7260], lang, 9.4)
    doc.add_page_break()


def contents_and_summary(doc, lang):
    zh = lang == "zh"
    add_heading(doc, "目录" if zh else "Contents", 1, lang)
    sections = [
        "1. 执行摘要" if zh else "1. Executive Summary",
        "2. 公司业务与海外目标客户" if zh else "2. Business Context and Overseas Target Buyers",
        "3. 海外获客平台规划" if zh else "3. Overseas Lead Source Strategy",
        "4. AI自动获客业务流程" if zh else "4. Automated Lead Generation Workflow",
        "5. 系统技术架构" if zh else "5. System Architecture",
        "6. 潜客评分与客户分层" if zh else "6. Lead Scoring and Qualification Tiers",
        "7. 公司数据需求" if zh else "7. Company Data Requirements",
        "8. 八周实施计划" if zh else "8. Eight-Week Implementation Plan",
        "9. 后续自动化路线" if zh else "9. Automation Roadmap",
        "10. 验收指标与风险控制" if zh else "10. Acceptance Metrics and Risk Controls",
        "附录" if zh else "Appendices",
    ]
    add_bullets(doc, sections, lang)
    add_heading(doc, "1. 执行摘要" if zh else "1. Executive Summary", 1, lang)
    if zh:
        add_para(doc, "本方案面向DPV International海外B2B业务增长需求，建设以多平台数据获取、企业识别、潜客评分、商品匹配、多语言触达和销售反馈为核心的自动获客系统。本期唯一核心客群为在目标国家从事商品进口和批发分销，并向当地连锁商店、区域零售网络及大型零售客户供货的海外进口商、批发商和分销商。B2C消费者获客不属于本期范围。", lang)
        add_para(doc, "首期项目周期为8周。第2周交付可演示原型，第4周完成多平台采集与统一潜客库MVP，第5周启动真实触达试运行，第8周完成验收和交接。第一阶段保留人工审批，确保客户信息、产品参数和外发内容均可追溯；运行稳定后，再逐步开放高分潜客自动首发及自动跟进。", lang)
        add_callout(doc, "项目核心成果：形成面向海外进口型批发商的“多平台发现—B2B资质验证—智能评分—商品匹配—人工审批—多渠道触达—销售跟进—成交反馈”业务闭环。", lang)
    else:
        add_para(doc, "This proposal establishes an overseas B2B lead generation system for DPV International, combining multi-platform company discovery, data enrichment, lead scoring, product-to-buyer matching, multilingual outreach and conversion feedback. The sole priority segment in the current phase is overseas importers, wholesalers and distributors that import merchandise and supply local chain stores, regional retail networks and large retail accounts. B2C customer acquisition is outside the scope of this phase.", lang)
        add_para(doc, "The initial delivery spans eight weeks. A demonstrable prototype will be available in Week 2, the multi-source lead database MVP will be completed in Week 4, live outreach will begin in Week 5 and acceptance will be completed in Week 8. Human approval remains mandatory during the initial phase, with controlled automation introduced only after data quality and operating performance are proven.", lang)
        add_callout(doc, "Primary outcome: a closed B2B operating loop for identifying, qualifying and converting overseas importer-wholesalers that supply chain-store and organized-retail channels.", lang)


def business_section(doc, lang):
    zh = lang == "zh"
    add_heading(doc, "2. 公司业务与海外目标客户" if zh else "2. Business Context and Overseas Target Buyers", 1, lang)
    if zh:
        add_para(doc, "DPV International对外定位为批量出口供应商，产品覆盖女装、男装、童装、箱包、美妆及家居日用品，并具备广州、香港和印度的供应与贸易协同基础。官网展示灵活MOQ、批发询价、FOB/CIF报价及出口包装等业务能力，可作为系统的企业知识库和商品知识来源。", lang)
        add_para(doc, "现有墨西哥市场用于提炼已验证的进口批发商画像、畅销产品和有效话术；阿联酋作为首个新市场试点，重点覆盖迪拜及周边具备进口能力、批发分销网络并向连锁商店供货的企业；孟加拉国作为第二拓展市场，依据历史数据和商品差异化能力复制同类B2B客户画像。", lang)
        add_heading(doc, "2.1 目标客户类型", 2, lang)
        add_bullets(doc, ["在目标国家依法开展商品进口、清关和批发分销的企业", "向连锁商店、区域零售网络、百货系统或大型零售客户持续供货的批发商与分销商", "具备可验证的进口记录、仓储或分销网络、稳定采购能力及持续补货需求", "采购负责人、进口经理、品类经理、批发业务负责人或企业决策人可被验证"], lang)
        add_heading(doc, "2.2 客户资格边界", 2, lang)
        add_para(doc, "本期仅处理符合上述画像的B2B企业。普通消费者、面向消费者的零售获客、仅经营单店或电商零售但不具备进口批发与连锁渠道供货能力的企业、自有品牌/OEM客户、采购代理，以及无企业主体的社交账号均不进入本期主动触达池。重复记录、已退订客户及无法验证来源的联系方式同样排除。", lang)
    else:
        add_para(doc, "DPV International is positioned as a bulk export supplier across women's, men's and children's apparel, bags, beauty products and household goods. Its operating footprint across Guangzhou, Hong Kong and India provides a foundation for sourcing and export coordination. The corporate website documents flexible MOQs, wholesale enquiries, FOB/CIF quotation practices and export-ready packing, and will serve as the system's company and product knowledge source.", lang)
        add_para(doc, "Mexico will provide the baseline for proven importer-wholesaler profiles, product demand and outreach language. The United Arab Emirates will be the first new-market pilot, focusing on Dubai-based companies with verifiable importing capability, wholesale distribution coverage and supply relationships with chain stores or organized retail networks. Bangladesh will follow as the second expansion market using the same B2B qualification model.", lang)
        add_heading(doc, "2.1 Target Buyer Segments", 2, lang)
        add_bullets(doc, ["Companies legally engaged in importing, customs clearance and wholesale distribution in the target market", "Wholesalers and distributors that supply chain stores, regional retail networks, department-store groups or other large retail accounts", "Businesses with verifiable import records, warehousing or distribution coverage, stable purchasing capacity and recurring replenishment demand", "Verifiable purchasing directors, import managers, category managers, wholesale managers or company decision-makers"], lang)
        add_heading(doc, "2.2 Qualification Boundaries", 2, lang)
        add_para(doc, "The current phase is limited to B2B companies that match the profile above. Individual consumers, consumer-retail acquisition, single-store or marketplace retailers without importer-wholesaler and chain-supply capability, private-label/OEM buyers, sourcing agencies and social accounts without a verifiable business entity are outside the active outreach pool. Duplicates, opted-out contacts and unverifiable contact records are also excluded.", lang)


def platform_section(doc, lang, fig):
    zh = lang == "zh"
    add_heading(doc, "3. 海外获客平台规划" if zh else "3. Overseas Lead Source Strategy", 1, lang)
    intro = "海外平台按照“进口批发商发现、连锁渠道供货证据、采购信号、联系人补全、B2B询盘承接”五类作用进行组合。社交媒体和内容平台仅用于发现企业与验证B2B信号，不用于获取普通消费者。" if zh else "Overseas platforms are combined across five roles: importer-wholesaler discovery, evidence of chain-store supply, buying signals, contact enrichment and B2B enquiry capture. Social and content platforms are used only for company discovery and B2B signal validation, not for consumer acquisition."
    add_para(doc, intro, lang)
    headers = ["平台类别", "代表渠道", "主要用途", "首期接入方式"] if zh else ["Source Type", "Representative Channels", "Primary Purpose", "Initial Integration"]
    rows_zh = [
        ("搜索与地图", "Google、Google Maps、Bing", "发现进口批发企业、仓储网点及地区覆盖", "搜索/API及公开页面解析"),
        ("职业与企业", "LinkedIn", "识别进口商、批发商、分销商及采购决策人", "研究任务与人工触达"),
        ("B2B社媒信号", "Facebook、Instagram", "发现企业主页、批发业务和连锁供货信号", "公开企业信号和自有账号数据"),
        ("行业内容信号", "TikTok、YouTube", "识别企业品牌、仓储分销和行业活动信号", "官方接口与公开元数据"),
        ("B2B采购平台", "Alibaba、Global Sources、TradeKey等", "发现进口采购商、批发询盘及品类需求", "平台授权、导出或API"),
        ("贸易与企业目录", "ImportGenius、Panjiva、Kompass等", "验证进口记录、批发身份、分销覆盖与采购能力", "公司账号或公开数据"),
        ("展会与供应商名单", "海外展会、行业协会、连锁渠道供应商名单", "发现参展批发商、进口分销商及连锁供货企业", "名单解析和企业主体匹配"),
    ]
    rows_en = [
        ("Search & Maps", "Google, Google Maps, Bing", "Identify importer-wholesalers, warehouses and geographic coverage", "Search/API and public-page parsing"),
        ("Professional Networks", "LinkedIn", "Identify importers, wholesalers, distributors and purchasing decision-makers", "Research tasks and human outreach"),
        ("B2B Social Signals", "Facebook, Instagram", "Discover company pages, wholesale activity and chain-supply evidence", "Public business signals and owned-account data"),
        ("Industry Content Signals", "TikTok, YouTube", "Identify corporate brands, warehousing, distribution and industry activity", "Official APIs and public metadata"),
        ("B2B Marketplaces", "Alibaba, Global Sources, TradeKey", "Discover import buyers, wholesale enquiries and category demand", "Platform-authorized export or API"),
        ("Trade Data & Directories", "ImportGenius, Panjiva, Kompass", "Validate import activity, wholesale status, distribution coverage and capacity", "Corporate accounts or public data"),
        ("Events & Supplier Lists", "Trade shows, trade associations, chain-supplier lists", "Find exhibiting wholesalers, import distributors and chain-store suppliers", "List parsing and entity matching"),
    ]
    add_table(doc, headers, rows_zh if zh else rows_en, [1500, 2200, 3400, 2260], lang, 7.7)
    add_figure(doc, fig, "图1 海外多平台客户获取来源图" if zh else "Figure 1. Multi-Platform Overseas Lead Sources",
               "平台数据首先汇总为统一企业记录，再执行清洗、去重、验证和资格判断。" if zh else "Platform signals are consolidated into unified company records before cleaning, deduplication, verification and qualification.", lang, 6.25)


def workflow_section(doc, lang, fig):
    zh = lang == "zh"
    add_heading(doc, "4. AI自动获客业务流程" if zh else "4. Automated Lead Generation Workflow", 1, lang)
    if zh:
        add_para(doc, "系统按国家、城市、商品、进口批发商画像和语言生成搜索策略，定时获取公开企业信号并形成候选企业。每条记录保留来源链接、采集时间和平台类型，经过标准化、去重、联系人验证、进口/批发资质判断、连锁渠道供货证据验证和商品匹配后进入评分。", lang)
        add_para(doc, "A级潜客生成个性化开发邮件、LinkedIn联系任务和产品推荐。首期所有外发内容经销售人员审批；客户回复、退订、退信或人工接管后，后续自动任务立即停止。回复被分类为询价、目录、样品、暂缓、拒绝或无关，并进入对应销售动作。", lang)
    else:
        add_para(doc, "The system generates search strategies by market, city, product, importer-wholesaler profile and language, then acquires public business signals on a scheduled basis. Every candidate record retains its source URL, capture date and platform type before standardization, deduplication, contact verification, importer/wholesaler qualification, chain-supply evidence validation and product matching.", lang)
        add_para(doc, "Tier A leads receive personalized email drafts, LinkedIn contact tasks and product recommendations. All outbound content requires sales approval during the initial phase. Follow-up stops immediately after a reply, opt-out, hard bounce or manual takeover. Replies are classified into quotation, catalogue, sample, defer, decline or irrelevant categories and routed to the appropriate sales action.", lang)
    add_figure(doc, fig, "图2 AI海外自动获客端到端流程图" if zh else "Figure 2. End-to-End AI-Powered Lead Generation Workflow",
               "流程包含无效客户淘汰、退订停止、高意向转销售及成交数据反馈四类控制。" if zh else "The workflow includes controls for invalid records, opt-outs, high-intent sales handoff and conversion feedback.", lang, 6.3)


def architecture_section(doc, lang, fig):
    zh = lang == "zh"
    add_heading(doc, "5. 系统技术架构" if zh else "5. System Architecture", 1, lang)
    if zh:
        add_para(doc, "系统采用“n8n工作流中枢 + PostgreSQL数据底座 + CRM业务界面 + 大模型能力服务”的分层架构。n8n负责任务调度、API连接、审批、告警和数据回写；PostgreSQL保存潜客、商品、活动、互动、退订和审计记录；已有CRM优先接入，未配置CRM时使用NocoDB提供销售操作界面。", lang)
        add_para(doc, "大模型用于企业分类、网页摘要、商品匹配、多语言文案和回复意图识别，不直接决定价格、付款条件、合同条款和交期承诺。所有密钥使用受控凭据存储，所有外发操作保留执行日志。", lang)
    else:
        add_para(doc, "The solution uses a layered architecture built around n8n orchestration, a PostgreSQL data foundation, a CRM operating interface and large-language-model services. n8n manages scheduling, API integrations, approvals, alerts and write-back. PostgreSQL stores leads, products, campaigns, interactions, opt-outs and audit records. An existing CRM will be integrated where available; otherwise, NocoDB will provide the initial sales interface.", lang)
        add_para(doc, "Language models support company classification, web-page summarization, product matching, multilingual copy and reply-intent detection. They do not independently determine prices, payment terms, contract terms or delivery commitments. Credentials are stored through controlled secrets management and all outbound actions are logged.", lang)
    add_landscape_figure(doc, "5.1 架构分层图" if zh else "5.1 Layered Architecture", fig,
                         "图3 AI自动获客系统技术架构图" if zh else "Figure 3. AI-Powered Lead Generation System Architecture",
                         "各层通过标准接口连接，可在不改变核心数据结构的情况下替换具体平台或模型服务。" if zh else "Standard interfaces allow individual platforms or model services to be replaced without changing the core data model.", lang)


def scoring_section(doc, lang, fig):
    zh = lang == "zh"
    add_heading(doc, "6. 潜客评分与客户分层" if zh else "6. Lead Scoring and Qualification Tiers", 1, lang)
    headers = ["评分维度", "权重", "判定依据"] if zh else ["Scoring Dimension", "Weight", "Evidence"]
    rows_zh = [("商品匹配", "20", "主营品类、价格带、MOQ和批发采购场景"), ("目标市场匹配", "15", "国家、城市及连锁零售覆盖区域"), ("进口批发商匹配", "15", "企业同时具备进口、批发或分销业务属性"), ("进口及连锁供货证据", "15", "贸易记录、仓储分销、批发说明及连锁渠道客户证据"), ("分销规模", "10", "仓库、员工、覆盖地区、营收或渠道网络信号"), ("近期采购信号", "10", "新品、招聘、参展、内容或进口活动"), ("采购负责人质量", "10", "职位、决策相关性及资料完整度"), ("联系方式有效性", "5", "工作邮箱、企业电话或可验证渠道")]
    rows_en = [("Product Fit", "20", "Category, price band, MOQ and wholesale buying use case"), ("Market Fit", "15", "Country, city and chain-retail coverage area"), ("Importer / Wholesaler Fit", "15", "Company has verifiable importing, wholesale or distribution functions"), ("Import & Chain-Supply Evidence", "15", "Trade records, warehousing, wholesale positioning and chain-customer evidence"), ("Distribution Scale", "10", "Warehouses, employees, territory, revenue or channel-network signals"), ("Recent Buying Signal", "10", "New ranges, recruitment, events, content or import activity"), ("Decision-Maker Quality", "10", "Role relevance and profile completeness"), ("Contact Validity", "5", "Verified work email, business phone or approved channel")]
    add_table(doc, headers, rows_zh if zh else rows_en, [2700, 1100, 5560], lang, 8.4)
    add_callout(doc, "A级：75–100分；B级：55–74分；C级：低于55分。85分以上仅在系统稳定且满足发送控制后进入自动首发候选。" if zh else "Tier A: 75–100; Tier B: 55–74; Tier C: below 55. Scores of 85 or above become eligible for automated first-touch only after the system is stable and send controls are satisfied.", lang, LIGHT_AMBER, AMBER)
    add_figure(doc, fig, "图4 潜客评分与客户分层图" if zh else "Figure 4. Lead Scoring and Qualification Tiers",
               "评分结果必须保留理由代码和来源证据，确保销售人员能够解释和复核。" if zh else "Every score retains reason codes and source evidence so sales teams can review and explain the result.", lang, 6.15)


def data_requirements(doc, lang):
    zh = lang == "zh"
    add_heading(doc, "7. 公司数据需求" if zh else "7. Company Data Requirements", 1, lang)
    add_para(doc, "HR承担跨部门协调职责，数据分别由管理层、销售、外贸、产品、供应链、财务和IT提供。首轮仅需要脱敏样本、字段说明和账号清单，不传递个人密码。" if zh else "HR coordinates the request across management, sales, export operations, product, supply chain, finance and IT. The first submission consists of de-identified samples, field definitions and an account inventory; individual passwords are not transferred.", lang)
    headers = ["优先级", "数据包", "核心内容", "范围", "责任部门"] if zh else ["Priority", "Data Package", "Required Content", "Coverage", "Owner"]
    rows_zh = [
        ("P0", "经营目标", "目标国家、进口批发商画像、连锁供货渠道、主推品类、线索/成交目标及排除范围", "管理层确认一页表", "管理层"),
        ("P0", "历史客户订单", "公司、国家、进口/批发属性、下游连锁渠道、SKU、数量、销售额、复购及周期", "最近12–24个月", "销售/财务"),
        ("P0", "历史线索漏斗", "来源、进口批发资质、连锁供货证据、联系人、触达、回复、报价、成交及未成交原因", "成功50家、失败100条以上", "销售/外贸"),
        ("P0", "商品主数据", "SKU、材料、尺寸、MOQ、阶梯价、FOB/CIF、样品、交期、产能、包装、认证和素材", "主推服装30–50个SKU", "产品/供应链"),
        ("P0", "销售知识", "公司介绍、FAQ、报价模板、成功话术、异议及禁止承诺", "成功案例20组以上", "销售/管理层"),
        ("P0", "系统与账号", "CRM、ERP、Excel、邮箱、WhatsApp、社媒、B2B平台、API和数据负责人", "完整系统盘点", "IT/行政"),
        ("P1", "渠道数据", "B2B官网询盘、WhatsApp企业询盘、展会、贸易目录、批发平台及企业社媒信号", "最近6–12个月", "市场/销售"),
        ("P1", "供应与履约", "广州、香港、印度的品类、产能、交期、港口、认证、投诉和限制", "按主体和品类", "供应链/工厂"),
        ("P1", "治理规则", "审批人、客户分配、退订名单、留存期限、发送上限和异常处理", "现行规则", "管理层/IT"),
    ]
    rows_en = [
        ("P0", "Business Objectives", "Target markets, importer-wholesaler ICP, chain-supply channels, priority categories, lead/conversion goals and exclusions", "One-page management brief", "Management"),
        ("P0", "Historical Customers & Orders", "Company, country, importer/wholesaler status, downstream chain channels, SKU, quantity, revenue, repeat orders and sales cycle", "Past 12–24 months", "Sales / Finance"),
        ("P0", "Historical Lead Funnel", "Source, importer-wholesaler qualification, chain-supply evidence, contact, outreach, reply, quotation, conversion and loss reason", "50 wins and 100+ losses/leads", "Sales / Export"),
        ("P0", "Product Master", "SKU, material, size, MOQ, price tiers, FOB/CIF, samples, lead time, capacity, packing, certification and assets", "30–50 priority apparel SKUs", "Product / Supply Chain"),
        ("P0", "Sales Knowledge", "Company profile, FAQs, quotation templates, successful copy, objections and prohibited claims", "20+ successful conversations", "Sales / Management"),
        ("P0", "Systems & Accounts", "CRM, ERP, Excel, email, WhatsApp, social platforms, B2B platforms, APIs and owners", "Complete inventory", "IT / Administration"),
        ("P1", "Channel Data", "B2B website and WhatsApp enquiries, trade shows, trade directories, wholesale platforms and corporate social signals", "Past 6–12 months", "Marketing / Sales"),
        ("P1", "Supply & Fulfilment", "Categories, capacity, lead times, ports, certifications, complaints and restrictions by operating location", "By entity and category", "Supply Chain / Factories"),
        ("P1", "Governance Rules", "Approvers, lead assignment, suppression list, retention, send limits and exception handling", "Current policies", "Management / IT"),
    ]
    add_table(doc, headers, rows_zh if zh else rows_en, [800, 1450, 4030, 1550, 1530], lang, 7.45)
    add_heading(doc, "7.1 首批样本量" if zh else "7.1 Initial Sample Size", 2, lang)
    add_bullets(doc, ["已成交进口商、批发商或分销商不少于50家；同类未成交线索不少于100条。" if zh else "At least 50 converted importers, wholesalers or distributors and 100 non-converted leads from the same segment.", "主推商品30–50个SKU；成功B2B沟通案例20组以上。" if zh else "30–50 priority SKUs and at least 20 successful B2B sales conversations.", "首批候选进口批发商200–500家；第一轮真实触达约100家。" if zh else "An initial pool of 200–500 candidate importer-wholesalers and approximately 100 live outreach records."], lang)


def implementation_plan(doc, lang, fig):
    zh = lang == "zh"
    if not zh:
        doc.add_page_break()
    add_heading(doc, "8. 八周实施计划" if zh else "8. Eight-Week Implementation Plan", 1, lang)
    headers = ["阶段", "时间", "主要工作", "交付成果", "验收点"] if zh else ["Phase", "Timing", "Primary Activities", "Deliverables", "Acceptance Gate"]
    rows_zh = [
        ("1 需求与数据盘点", "第1周", "访谈、数据样本、产品和市场确认", "数据清单、画像、系统盘点", "目标与字段确认"),
        ("2 快速原型", "第2周", "数据库、n8n、首个搜索流程", "50条样例潜客", "原型演示"),
        ("3 多平台采集MVP", "第3–4周", "多来源获取、清洗、去重与进口批发资质验证", "200–500家进口批发商库", "MVP完成"),
        ("4 分类与商品匹配", "第4–5周", "连锁供货证据、评分、商品匹配和多语言话术", "A/B/C分层及草稿", "抽查准确率"),
        ("5 审批与真实触达", "第5周", "人工审批、发送、回复监控和停止规则", "约100家真实触达", "试运行启动"),
        ("6 B2B信号与看板", "第6周", "企业社媒信号、B2B询盘归集和渠道报表", "管理看板与异常提醒", "数据可追溯"),
        ("7 优化与市场复制", "第7周", "根据结果优化并复制第二市场", "孟加拉国工作流", "规则复核"),
        ("8 验收与交接", "第8周", "测试、培训、文档和权限检查", "系统、手册、验收报告", "首期验收"),
    ]
    rows_en = [
        ("1 Requirements & Data", "Week 1", "Interviews, data samples, product and market confirmation", "Data register, ICP and system inventory", "Scope and fields approved"),
        ("2 Rapid Prototype", "Week 2", "Database, n8n and first search workflow", "50 sample leads", "Prototype demonstration"),
        ("3 Multi-Source MVP", "Weeks 3–4", "Acquisition, cleaning, deduplication and importer-wholesaler verification", "Pool of 200–500 importer-wholesalers", "MVP completed"),
        ("4 Qualification & Matching", "Weeks 4–5", "Chain-supply evidence, scoring, product matching and multilingual copy", "A/B/C tiers and drafts", "Accuracy sample passed"),
        ("5 Approval & Live Outreach", "Week 5", "Human approval, sending, reply monitoring and stop rules", "Approximately 100 live contacts", "Pilot launched"),
        ("6 B2B Signals & Dashboard", "Week 6", "Corporate social signals, B2B enquiry intake and channel reporting", "Management dashboard and alerts", "Full traceability"),
        ("7 Optimize & Replicate", "Week 7", "Optimize from pilot results and copy to second market", "Bangladesh workflow", "Rules reviewed"),
        ("8 Acceptance & Handover", "Week 8", "Testing, training, documentation and access review", "System, guide and acceptance report", "Initial acceptance"),
    ]
    add_table(doc, headers, rows_zh if zh else rows_en, [1700, 1000, 3300, 2100, 1260], lang, 7.4)
    add_landscape_figure(doc, "8.1 项目进度图" if zh else "8.1 Project Roadmap", fig,
                         "图5 八周项目实施进度图" if zh else "Figure 5. Eight-Week Implementation Roadmap",
                         "里程碑使管理层在第2、4、5和8周获得可验证成果，而非等待项目末期集中交付。" if zh else "Milestones provide management with verifiable outcomes in Weeks 2, 4, 5 and 8 instead of deferring all value to project close.", lang)


def roadmap_and_controls(doc, lang):
    zh = lang == "zh"
    add_heading(doc, "9. 后续自动化路线" if zh else "9. Automation Roadmap", 1, lang)
    headers = ["阶段", "时间", "自动化范围", "保留控制"] if zh else ["Stage", "Timing", "Automation Scope", "Controls Retained"]
    rows_zh = [("首期B2B", "8周内", "进口批发商搜索、资质验证、评分、匹配和话术生成", "所有外发人工审批；价格、合同和交期人工确认"), ("第二期B2B", "验收后4–6周", "85分以上进口批发商自动首发、定时跟进、目录发送和回复分流", "每日上限、退订、退信及异常停止"), ("第三期B2B", "稳定运行后6–8周", "扩展国家、品类、B2B询盘和评分优化", "报价、付款、合同及关键承诺继续审批"), ("未来B2C", "另行立项", "消费者获客、零售营销及相关渠道自动化", "不属于本期范围，需单独审批、数据和实施方案")]
    rows_en = [("Initial B2B", "Within 8 weeks", "Importer-wholesaler search, qualification, scoring, matching and outreach drafting", "Human approval for all sends; human confirmation for price, contract and delivery"), ("Stage 2 B2B", "4–6 weeks after acceptance", "Automated first-touch for 85+ importer-wholesalers, follow-up, catalogue delivery and reply routing", "Daily caps, opt-out, bounce and exception stops"), ("Stage 3 B2B", "6–8 weeks after stable operation", "Additional markets, categories, B2B enquiries and scoring optimization", "Approval remains for quotations, payment, contracts and material commitments"), ("Future B2C", "Separate project", "Consumer acquisition, retail marketing and related channel automation", "Outside the current scope; requires separate approval, data and implementation plan")]
    add_table(doc, headers, rows_zh if zh else rows_en, [1300, 1600, 3900, 2560], lang, 8.1)
    add_heading(doc, "10. 验收指标与风险控制" if zh else "10. Acceptance Metrics and Risk Controls", 1, lang)
    metrics_zh = [("字段完整率", "≥90%"), ("重复率", "<2%"), ("A级进口批发商抽查准确率", "≥80%"), ("硬退信率", "<5%"), ("来源可追溯", "100%"), ("首期外发审批", "100%"), ("人工找客时间降低", "≥60%")]
    metrics_en = [("Required-field completeness", "≥90%"), ("Duplicate rate", "<2%"), ("Tier A importer-wholesaler precision", "≥80%"), ("Hard-bounce rate", "<5%"), ("Source traceability", "100%"), ("Initial-phase send approval", "100%"), ("Manual research time reduction", "≥60%")]
    add_table(doc, ["指标" if zh else "Metric", "目标" if zh else "Target"], metrics_zh if zh else metrics_en, [6000, 3360], lang, 9)
    add_heading(doc, "10.1 关键风险控制" if zh else "10.1 Key Risk Controls", 2, lang)
    if zh:
        add_bullets(doc, ["范围控制：仅允许符合进口批发商画像且具有连锁渠道供货证据的企业进入触达；消费者、普通零售商及其他非目标B2B类型自动排除。", "平台限制：优先采用官方接口、授权导出和公开企业数据，平台不支持自动私信时生成销售任务。", "数据质量：保留来源、采集时间、验证状态和评分理由，重复及过期记录不进入触达。", "内容准确：文案仅引用商品主数据和已批准企业知识，缺失参数不自动补全。", "发送治理：执行退订名单、硬退信停止、每日上限、域名信誉监测和人工接管。", "商业承诺：报价、付款、合同、认证和交期承诺始终需要业务人员确认。", "权限安全：凭据分级、最小权限、操作日志和定期数据留存审查。"], lang, size=9.5, after=2)
    else:
        add_bullets(doc, ["Scope control: only companies matching the importer-wholesaler ICP with evidence of chain-store supply may enter outreach; consumers, ordinary retailers and other non-target B2B types are automatically excluded.", "Platform constraints: prioritize official APIs, authorized exports and public business data; create human sales tasks where automated messaging is unavailable.", "Data quality: retain source, capture time, verification status and score reasons; duplicates and stale records do not enter outreach.", "Content accuracy: outbound copy references only approved company knowledge and product master data; missing specifications are not invented.", "Send governance: enforce suppression lists, hard-bounce stops, daily caps, domain-reputation monitoring and manual takeover.", "Commercial commitments: quotations, payment, contracts, certifications and delivery commitments always require business confirmation.", "Access security: apply credential segregation, least privilege, activity logging and periodic retention review."], lang, size=9.5, after=2)


def appendices(doc, lang):
    zh = lang == "zh"
    add_heading(doc, "附录A：HR数据协调清单" if zh else "Appendix A: HR Data Coordination Checklist", 1, lang)
    checks_zh = ["安排管理层、销售、外贸、产品、供应链、财务和IT数据负责人", "提供Customers_Orders、Leads_Funnel、Product_Master、Systems_Channels四类Excel样本", "提供Sales_Messages_FAQ文档及产品图片/视频文件夹", "确认字段定义、币种、日期格式、空值含义和数据更新频率", "建立企业账号授权流程，不通过表格传递个人密码", "确认数据留存、退订、审批和客户分配负责人"]
    checks_en = ["Nominate data owners across management, sales, export operations, product, supply chain, finance and IT", "Provide sample Customers_Orders, Leads_Funnel, Product_Master and Systems_Channels workbooks", "Provide Sales_Messages_FAQ documentation and product image/video folders", "Confirm field definitions, currencies, date formats, null-value meanings and update frequency", "Establish corporate account authorization without transferring personal passwords in spreadsheets", "Confirm owners for retention, opt-out, approval and lead assignment"]
    add_bullets(doc, checks_zh if zh else checks_en, lang, size=9.6, after=2)
    add_heading(doc, "附录B：核心数据字段" if zh else "Appendix B: Core Data Fields", 1, lang)
    rows_zh = [("LeadRecord", "公司、国家、网站、来源、进口/批发属性、连锁供货证据、分销区域、采购联系人、工作邮箱、语言、评分、阶段、负责人、退订"), ("ProductRecord", "SKU、品类、材质、尺寸、MOQ、阶梯价、贸易条款、样品、交期、产能、包装、认证、素材"), ("InteractionRecord", "潜客ID、渠道、时间、方向、模板、活动、打开、点击、回复、退信、意图和处理结果"), ("CampaignRecord", "市场、进口批发商画像、商品、渠道、关键词、话术版本、发送上限、线索、回复、报价和成交")]
    rows_en = [("LeadRecord", "Company, country, website, source, importer/wholesaler status, chain-supply evidence, distribution territory, purchasing contact, work email, language, score, stage, owner and opt-out"), ("ProductRecord", "SKU, category, material, size, MOQ, price tier, Incoterm, sample, lead time, capacity, packing, certification and assets"), ("InteractionRecord", "Lead ID, channel, time, direction, template, campaign, opens, clicks, replies, bounces, intent and disposition"), ("CampaignRecord", "Market, importer-wholesaler ICP, products, channel, keywords, copy version, send cap, leads, replies, quotations and conversions")]
    add_table(doc, ["数据表", "核心字段"] if zh else ["Record", "Core Fields"], rows_zh if zh else rows_en, [2200, 7160], lang, 8.1)
    add_heading(doc, "附录C：参考资料" if zh else "Appendix C: References", 1, lang)
    refs = [
        "DPV International. Corporate website and product catalogue. https://dpvinternational.com/",
        "n8n Documentation. Workflow automation and AI integrations. https://docs.n8n.io/",
        "Google for Developers. YouTube Data API. https://developers.google.com/youtube/v3/docs",
        "LinkedIn Help. Accessing LinkedIn APIs. https://www.linkedin.com/help/linkedin/answer/a526048",
        "Meta. WhatsApp Cloud API documentation.",
    ]
    add_bullets(doc, refs, lang, size=7.8, after=1)


def build_document(lang, figures, logo):
    doc = Document()
    configure_section(doc.sections[0], False)
    setup_styles(doc, lang)
    cover(doc, lang, logo)
    contents_and_summary(doc, lang)
    business_section(doc, lang)
    platform_section(doc, lang, figures["source"])
    workflow_section(doc, lang, figures["flow"])
    architecture_section(doc, lang, figures["arch"])
    scoring_section(doc, lang, figures["score"])
    data_requirements(doc, lang)
    implementation_plan(doc, lang, figures["timeline"])
    roadmap_and_controls(doc, lang)
    appendices(doc, lang)
    set_headers_footers(doc, lang)
    doc.core_properties.title = "DPV International AI海外自动获客系统建设方案" if lang == "zh" else "DPV International AI-Powered Overseas Lead Generation System Proposal"
    doc.core_properties.subject = "Overseas lead generation system implementation proposal"
    doc.core_properties.author = "DPV International"
    doc.core_properties.last_modified_by = "DPV International"
    doc.core_properties.comments = ""
    path = OUT / ("DPV International AI海外自动获客系统建设方案（中文版）.docx" if lang == "zh" else "DPV International AI-Powered Overseas Lead Generation System Proposal（English）.docx")
    doc.save(path)
    return path


def main():
    logo = download_logo()
    all_figs = {}
    for lang in ("zh", "en"):
        all_figs[lang] = {
            "source": source_map(lang),
            "flow": acquisition_flow(lang),
            "arch": architecture(lang),
            "score": scoring(lang),
            "timeline": timeline(lang),
        }
    paths = [build_document("zh", all_figs["zh"], logo), build_document("en", all_figs["en"], logo)]
    for p in paths:
        print(p)


if __name__ == "__main__":
    main()
